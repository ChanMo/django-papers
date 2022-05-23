from urllib.parse import urlparse
from docx import Document
from django.conf import settings

def get_img_path(url):
    src = urlparse(url)
    src = settings.BASE_DIR / src.path[1:]
    return str(src)


def draft_to_docx(data, outfile):
    """
    Draft.js contentState to docx

    TODO: use cachefile
    TODO: heading inlinestyle and entityranges
    TODO: uncollapse entity

    Parameter
    ----------
    data: JSON
        draft.js contentState
    outfile: String
        output filename

    Response
    ----------
    string: file path
    """
    d = Document()
    for block in data['blocks']:
        if block['type'] == 'header-one':
            d.add_heading(block['text'], 1)
        elif block['type'] == 'header-two':
            d.add_heading(block['text'], 2)
        elif block['type'] == 'header-three':
            d.add_heading(block['text'], 3)
        elif block['type'] == 'header-four':
            d.add_heading(block['text'], 4)
        elif block['type'] == 'unstyled':
            p = d.add_paragraph()
            for index, text in enumerate(block['text']):
                if index in [i['offset'] for i in block['entityRanges']]:
                    for entity_row in block['entityRanges']:
                        if entity_row['offset'] == index:
                            entity = data['entityMap'].get(str(entity_row['key']))
                            if entity['type'] == 'TEX':
                                # for test
                                p.add_run('TEX_DEMO')
                            elif entity['type'] == 'IMG':
                                p.add_run().add_picture(get_img_path(entity['data']['src']))
                            break

                run = p.add_run(text)
                for inline in block['inlineStyleRanges']:
                    if inline['offset'] <= index <= inline['offset'] + inline['length']:
                        if inline['style'] == 'BOLD':
                            run.font.bold = True
                        elif inline['style'] == 'ITALIC':
                            run.font.italic = True
                        elif inline['style'] == 'UNDERLINE':
                            run.font.underline = True
                        elif inline['style'] == 'SUPERSCRIPT':
                            run.font.superscript = True
                        elif inline['style'] == 'SUBSCRIPT':
                            run.font.subscript = True



        elif block['type'] == 'atomic':
            pass


    res = f'/tmp/{outfile}.docx'
    d.save(res)
    return res
